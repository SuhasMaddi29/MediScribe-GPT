from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
import anthropic
import google.generativeai as genai
import os
import json
import tempfile
import assemblyai as aai
import time

app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Adjust origins as needed
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"]
)

# Paths for the Word template and reports directory
TEMPLATE_PATH = "report.docx"  # Template file path
REPORTS_DIR = os.path.join(os.getcwd(), "tmp")  # Reports directory path
os.makedirs(REPORTS_DIR, exist_ok=True)  # Ensure the directory exists
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

def process_with_lmstudio(transcribed_text):
    try:
        print("Inside LMStudio processing")
        
        # Define the payload
        payload = {
            "model": "qwen2.5-14b-instruct",
            "messages": [{"role": "system", "content": "The following text is a transcription of a meeting. Extract key information "
                "and return it in the form of a JSON dictionary. The keys should include:\n"
                " - 'doctor' (Doctor's name)\n"
                " - 'visit_date' (Date of visit to doctor)\n"
                " - 'specialization' (Specialization of the doctor)\n"
                " - 'patient' (Patient's name)\n"
                " - 'birth_date' (Patient's Date of birth)\n"
                " - 'medNumber' (Patient's medical number)\n"
                " - 'ihi' (Patient's Individual Healthcare Identifier)\n"
                " - 'patientPhone' (Patient's Phone number)\n"
                " - 'email' (Patient's email)\n"
                " - 'medical_history' (Medical History of the patient) (GIVE THIS IN BULLET POINT FORMAT)\n"
                " - 'assessment' (Current assessment of the patient from Doctor's POV) (GIVE THIS IN BULLET POINT FORMAT)\n"
                " - 'diagnosis' (Diagnosis of the disease) (If there is a suggested diagnosis that you can interpret, mention this as suggestion by AI)\n"
                " - 'prescription' (Medications prescribed by the doctor) (If there is a suggested prescription that you can interpret, mention this as suggestion by AI)\n\n"
                f"Here is the transcription:\n\n{transcribed_text}\n\n"
                "Return only the JSON object without any additional text or explanation.\n\n"
                "If a key is not available, return as Unknown\n"}],
            "temperature": 0.7,
            "max_tokens": 4096,
            "stream": False
        }

        # Make the HTTP POST request to the LMStudio API
        response = requests.post(
            "http://localhost:1234/api/v0/chat/completions",
            headers={"Content-Type": "application/json"},
            json=payload
        )

        # Raise an error if the request failed
        response.raise_for_status()

        # Parse and clean the response
        print("Raw Response:", response.text)
        json_response = response.json()
        content = json_response["choices"][0]["message"]["content"]
        try:
            json_response = json.loads(content)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format: {e}")

        return json_response

    except requests.RequestException as e:
        raise RuntimeError(f"Error communicating with LMStudio API: {str(e)}")
    except Exception as e:
        raise RuntimeError(f"Error processing text with LMStudio: {str(e)}")


def process_with_gemini(transcribed_text):
    print("Inside Gemini processing")

    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel("gemini-1.5-pro")
    try:
        response = model.generate_content(f"The following text is a transcription of a meeting. Extract key information "
                "and return it in the form of a JSON dictionary. The keys should include:\n"
                " - 'doctor' (Doctor's name)\n"
            " - 'visit_date' (Date of visit to doctor)\n"
            " - 'specialization' (Specialization of the doctor)\n"
            " - 'patient' (Patient's name)\n"
            " - 'birth_date' (Patient's Date of birth)\n"
            " - 'medNumber' (Patient's medical number)\n"
            " - 'ihi' (Patient's Individual Healthcare Identifier)\n"
            " - 'patientPhone' (Patient's Phone number)\n"
            " - 'email' (Patient's email)\n"
            " - 'medical_history' (Medical History of the patient) (GIVE THIS IN BULLET POINT FORMAT)\n"
            " - 'assessment' (Current assessment of the patient from Doctor's POV) (GIVE THIS IN BULLET POINT FORMAT)\n"
            " - 'diagnosis' (Diagnosis of the disease) (If there is a suggested diagnosis that you can interpret, mention this as suggestion by AI)\n"
            " - 'prescription' (Medications prescribed by the doctor) (If there is a suggested prescription that you can interpret, mention this as suggestion by AI)\n\n"
            f"Here is the transcription:\n\n{transcribed_text}\n\n"
            "Return only the JSON object without any additional text or explanation.\n\n"
            "If a key is not available, return as Unknown\n"
            "Please return only the object with braces and not other additional text like json so that i can directly process it using json loads")

        print("Raw Response:", response.text)

        try:
            cleaned_input = response.text.strip("```json").rstrip("```").strip()
            cleaned_input = cleaned_input.replace('```', '')
            # Parse as JSON
            return json.loads(cleaned_input)
        
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format: {e}")

    except Exception as e:
        raise RuntimeError(f"Error processing text with Gemini: {str(e)}")
    
# Claude integration
def process_with_claude(transcribed_text):
    client = anthropic.Anthropic()
    try:
        print("Inside Claude processing")

        # Define the prompt
        prompt = (
            f"{anthropic.HUMAN_PROMPT} The following text is a transcription of a meeting. Extract key information "
            "and return it in the form of a JSON dictionary. The keys should include:\n"
            " - 'doctor' (Doctor's name)\n"
            " - 'visit_date' (Date of visit to doctor)\n"
            " - 'specialization' (Specialization of the doctor)\n"
            " - 'patient' (Patient's name)\n"
            " - 'birth_date' (Patient's Date of birth)\n"
            " - 'medNumber' (Patient's medical number)\n"
            " - 'ihi' (Patient's Individual Healthcare Identifier)\n"
            " - 'patientPhone' (Patient's Phone number)\n"
            " - 'email' (Patient's email)\n"
            " - 'medical_history' (Medical History of the patient) (GIVE THIS IN BULLET POINT FORMAT)\n"
            " - 'assessment' (Current assessment of the patient from Doctor's POV) (GIVE THIS IN BULLET POINT FORMAT)\n"
            " - 'diagnosis' (Diagnosis of the disease) (If there is a suggested diagnosis that you can interpret, mention this as suggestion by AI)\n"
            " - 'prescription' (Medications prescribed by the doctor) (If there is a suggested prescription that you can interpret, mention this as suggestion by AI)\n\n"
            f"Here is the transcription:\n\n{transcribed_text}\n\n"
            "Return only the JSON object without any additional text or explanation.\n\n"
            "If a key is not available, return as Unknown\n"
            f"{anthropic.AI_PROMPT}"
        )

        # Call the Claude API
        message = client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=1024,
            messages=[
                {"role": "user", "content": prompt}
            ]
        ).content[0].text

        print("Raw Response:", message)

        # Directly parse the JSON string from the response
        try:
            json_response = json.loads(message)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON returned by Claude: {str(e)}")

        return json_response

    except Exception as e:
        raise RuntimeError(f"Error processing text with Claude: {str(e)}")


# Function to populate the Word document
def populate_docx(template_path, output_path, data):
    try:
        doc = Document(template_path)

        def format_list_as_bullets(items):
            return '\n'.join([f"• {item}" for item in items])
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"  # Placeholder format: {{key}}
                if isinstance(value, list):
                    # Join list items with newlines or bullets
                    value = format_list_as_bullets(value)
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if isinstance(value, list):
                            # Join list items with newlines or bullets
                            value = format_list_as_bullets(value)
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
        
        # Save the modified document
        doc.save(output_path)
        print(f"Document saved at {output_path}")
    except Exception as e:
        print(f"Error populating Word document: {str(e)}")
        raise


def flatten_data(data, parent_key='', sep='.'):
    items = []
    for k, v in data.items():
        new_key = f"{parent_key}{sep}{k}" if parent_key else k
        if isinstance(v, dict):
            items.extend(flatten_data(v, new_key, sep=sep).items())
        else:
            items.append((new_key, v))
    return dict(items)


@app.post("/process-audio")
async def process_audio(audio_file: UploadFile, ai_model: str = Form(...)):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp_audio:
            temp_audio.write(audio_file.file.read())
            temp_audio_path = temp_audio.name

        uploaded_file_name = audio_file.filename
        audio_file_base = os.path.splitext(uploaded_file_name)[0]  # Strip extension
        timestamp = time.strftime("%Y%m%d%H%M")  # Current timestamp

        OUTPUT_PATH = os.path.join(
            os.getcwd(),
            "tmp",
            f"{audio_file_base}_{ai_model}_{timestamp}.docx"
        )
        
        os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)

        audio_file = (
                temp_audio_path
            )

        config = aai.TranscriptionConfig(
            speaker_labels=True,
        )

        transcript = aai.Transcriber().transcribe(audio_file, config)
        transcribed_text = ""

        for utterance in transcript.utterances:
            if utterance.speaker == "A":
                transcribed_text += "Doctor: " + utterance.text + " "
            else:
                transcribed_text += "Patient: " + utterance.text + " "

        print(transcribed_text)

        # Step 4: Use ChatGPT to extract key-value pairs
        if ai_model=="gemini":
            key_value_pairs = process_with_gemini(transcribed_text)
        elif ai_model == "claude":
            key_value_pairs = process_with_claude(transcribed_text)
        else:
            key_value_pairs = process_with_lmstudio(transcribed_text)

        key_value_pairs = flatten_data(key_value_pairs)

        populate_docx(TEMPLATE_PATH, OUTPUT_PATH, key_value_pairs)
        print(f"Document saved at {OUTPUT_PATH}")
        # Comment out the send_file line for now
        # return send_file(OUTPUT_PATH, as_attachment=True)
        # return {"message": f"File saved at {OUTPUT_PATH}"}

        response = FileResponse(
                    OUTPUT_PATH,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={
                        "Content-Disposition": f'attachment; filename="Report_{audio_file_base}_{timestamp}.docx"'
                    },
                )
        print(f"Response Headers: {response.headers}")
        report_metadata = {
            "filename": f"Report_{audio_file_base}_{timestamp}.docx",
            "ai_model": ai_model,
            "timestamp": timestamp
        }

        # Path to the metadata file
        metadata_file_path = os.path.join(REPORTS_DIR, "reports_metadata.json")

        # Load existing metadata if the file exists
        if os.path.exists(metadata_file_path):
            with open(metadata_file_path, "r") as f:
                metadata = json.load(f)
        else:
            metadata = []

        # Append the new report metadata
        metadata.append(report_metadata)

        # Save the updated metadata
        with open(metadata_file_path, "w") as f:
            json.dump(metadata, f, indent=4)

        # Return a success message
        return response
    
    except Exception as e:
        return {'error': str(e)}, 500
    finally:
        if os.path.exists(temp_audio_path):
            os.remove(temp_audio_path)



@app.get("/list-reports")
async def list_reports():
    try:
        # Ensure the reports directory exists
        if not os.path.exists(REPORTS_DIR):
            return JSONResponse(content={"reports": []})

        # List all .docx files in the reports directory
        files = []
        for filename in os.listdir(REPORTS_DIR):
            if filename.endswith(".docx"):
                file_path = os.path.join(REPORTS_DIR, filename)
                if os.path.isfile(file_path):
                    # Get the file's modification time
                    modified_time = os.path.getmtime(file_path)
                    files.append({
                        "name": filename,
                        "modified_time": modified_time
                    })

        # Sort files by modification time (most recent first)
        files.sort(key=lambda x: x["modified_time"], reverse=True)

        # Return the list of files
        return JSONResponse(content={"reports": files})
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)

@app.get("/download-report/{filename}")
async def download_report(filename: str):
    try:
        # Sanitize the filename to prevent directory traversal
        safe_filename = os.path.basename(filename)
        file_path = os.path.join(REPORTS_DIR, safe_filename)

        # Check if the file exists
        if not os.path.isfile(file_path):
            return JSONResponse(content={"error": "File not found"}, status_code=404)

        # Return the file
        return FileResponse(
            path=file_path,
            filename=safe_filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)